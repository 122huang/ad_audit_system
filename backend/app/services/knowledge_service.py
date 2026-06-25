import os
import re
from typing import List, Optional
from datetime import datetime
from sqlalchemy.orm import Session
from sqlalchemy import or_
import logging
import uuid

from app.models.models import (
    KnowledgeDocument, DocumentChunk, AuditCase,
    CustomRule, Region, DocType
)
from app.core.config import settings

logger = logging.getLogger(__name__)


class DocumentParser:
    """文档解析器 - 支持多格式"""

    @staticmethod
    def parse_txt(content: bytes) -> str:
        return content.decode('utf-8', errors='ignore')

    @staticmethod
    def parse_pdf(content: bytes) -> str:
        try:
            from io import BytesIO
            from pypdf import PdfReader
            reader = PdfReader(BytesIO(content))
            texts = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                texts.append(f"[Page {i+1}]\n{text}")
            return "\n\n".join(texts)
        except Exception as e:
            logger.error(f"PDF parse error: {e}")
            return ""

    @staticmethod
    def parse_docx(content: bytes) -> str:
        try:
            from io import BytesIO
            from docx import Document
            doc = Document(BytesIO(content))
            texts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    texts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    texts.append(row_text)
            return "\n".join(texts)
        except Exception as e:
            logger.error(f"DOCX parse error: {e}")
            return ""

    @staticmethod
    def parse_excel(content: bytes) -> str:
        try:
            from io import BytesIO
            from openpyxl import load_workbook
            wb = load_workbook(BytesIO(content), read_only=True)
            texts = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                texts.append(f"[Sheet: {sheet_name}]")
                for row in ws.iter_rows(values_only=True):
                    row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    if row_text.strip(" |"):
                        texts.append(row_text)
            return "\n".join(texts)
        except Exception as e:
            logger.error(f"Excel parse error: {e}")
            return ""


class TextChunker:
    """智能文本分块器 - 保留标题层级和列表结构"""

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def split_text(self, text: str) -> List[dict]:
        chunks = []

        lines = text.split('\n')
        current_section = ""
        current_title = ""
        chunk_index = 0

        buffer = ""

        for line in lines:
            heading_match = re.match(r'^(#{1,6}\s+|第[一二三四五六七八九十\d]+[章节条款]\s*|\d+\.\d+\s+|\d+\.\s+)', line.strip())
            list_match = re.match(r'^(\s*)([-*•]|\d+[.)、])\s+', line)

            if heading_match and len(buffer) > 100:
                if buffer.strip():
                    chunks.append({
                        "index": chunk_index,
                        "content": buffer.strip(),
                        "metadata": {"title": current_title, "type": "section"}
                    })
                    chunk_index += 1
                current_title = line.strip()
                buffer = line + "\n"
            elif list_match and buffer and not buffer.endswith('\n'):
                if len(buffer) > self.chunk_size:
                    chunks.append({
                        "index": chunk_index,
                        "content": buffer.strip(),
                        "metadata": {"title": current_title, "type": "list"}
                    })
                    chunk_index += 1
                    buffer = line + "\n"
                else:
                    buffer += line + "\n"
            else:
                buffer += line + "\n"

                if len(buffer) >= self.chunk_size:
                    chunks.append({
                        "index": chunk_index,
                        "content": buffer.strip(),
                        "metadata": {"title": current_title, "type": "content"}
                    })
                    chunk_index += 1
                    overlap_text = buffer[-self.chunk_overlap:] if self.chunk_overlap > 0 else ""
                    buffer = overlap_text

        if buffer.strip():
            chunks.append({
                "index": chunk_index,
                "content": buffer.strip(),
                "metadata": {"title": current_title, "type": "content"}
            })

        return chunks


class KnowledgeService:
    """知识库管理服务"""

    def __init__(self, db: Session):
        self.db = db
        self.chunker = TextChunker(chunk_size=500, chunk_overlap=50)
        self.parser = DocumentParser()
        self.upload_dir = settings.UPLOAD_DIR
        os.makedirs(self.upload_dir, exist_ok=True)

    def create_document_from_text(self, title: str, doc_type: str, content: str,
                                   region_code: Optional[str] = None,
                                   tags: Optional[List[str]] = None,
                                   source: Optional[str] = None,
                                   created_by: str = "system") -> KnowledgeDocument:
        doc = KnowledgeDocument(
            id=str(uuid.uuid4()),
            title=title,
            doc_type=DocType[doc_type] if doc_type in DocType.__members__ else DocType.internal,
            content_text=content,
            source=source,
            tags=tags or [],
            created_by=created_by,
            is_indexed=False
        )

        if region_code:
            region = self.db.query(Region).filter(Region.code == region_code).first()
            if region:
                doc.region_id = region.id

        self.db.add(doc)
        self.db.commit()
        self.db.refresh(doc)

        self._index_document(doc)
        return doc

    def upload_document(self, file_content: bytes, filename: str,
                        doc_type: str, title: Optional[str] = None,
                        region_code: Optional[str] = None,
                        tags: Optional[List[str]] = None,
                        created_by: str = "system") -> KnowledgeDocument:
        ext = os.path.splitext(filename)[1].lower()

        if ext == '.pdf':
            content = self.parser.parse_pdf(file_content)
        elif ext in ['.docx', '.doc']:
            content = self.parser.parse_docx(file_content)
        elif ext in ['.xlsx', '.xls']:
            content = self.parser.parse_excel(file_content)
        elif ext in ['.txt', '.md']:
            content = self.parser.parse_txt(file_content)
        else:
            content = file_content.decode('utf-8', errors='ignore')

        safe_filename = f"{uuid.uuid4().hex}{ext}"
        file_path = os.path.join(self.upload_dir, safe_filename)
        with open(file_path, 'wb') as f:
            f.write(file_content)

        doc_title = title or os.path.splitext(filename)[0]

        doc = KnowledgeDocument(
            id=str(uuid.uuid4()),
            title=doc_title,
            doc_type=DocType[doc_type] if doc_type in DocType.__members__ else DocType.internal,
            file_name=filename,
            file_url=f"/uploads/{safe_filename}",
            content_text=content,
            tags=tags or [],
            created_by=created_by,
            is_indexed=False
        )

        if region_code:
            region = self.db.query(Region).filter(Region.code == region_code).first()
            if region:
                doc.region_id = region.id

        self.db.add(doc)
        self.db.commit()
        self.db.refresh(doc)

        self._index_document(doc)
        return doc

    def _index_document(self, doc: KnowledgeDocument):
        try:
            chunks = self.chunker.split_text(doc.content_text or "")

            for chunk_data in chunks:
                chunk = DocumentChunk(
                    id=str(uuid.uuid4()),
                    document_id=doc.id,
                    chunk_index=chunk_data["index"],
                    content=chunk_data["content"],
                    chunk_metadata=chunk_data["metadata"]
                )
                self.db.add(chunk)

            doc.chunk_count = len(chunks)
            doc.is_indexed = True
            self.db.commit()
            logger.info(f"Document {doc.id} indexed with {len(chunks)} chunks")
        except Exception as e:
            logger.error(f"Index error for doc {doc.id}: {e}")
            self.db.rollback()

    def create_custom_rule(self, user_id: str, rule_name: str, rule_type: str,
                           keywords: List[str], patterns: List[str],
                           description: str, suggestion: Optional[str] = None,
                           region_code: Optional[str] = None,
                           severity: str = "medium") -> CustomRule:
        rule_code = f"CUS_{uuid.uuid4().hex[:8].upper()}"

        rule = CustomRule(
            id=str(uuid.uuid4()),
            user_id=user_id,
            rule_code=rule_code,
            rule_name=rule_name,
            rule_type=rule_type,
            severity=severity,
            keywords=keywords,
            patterns=patterns,
            description=description,
            suggestion=suggestion,
            is_active=True
        )

        if region_code:
            region = self.db.query(Region).filter(Region.code == region_code).first()
            if region:
                rule.region_id = region.id

        self.db.add(rule)
        self.db.commit()
        self.db.refresh(rule)
        return rule

    def create_case(self, title: str, material_type: str, violation_type: str,
                    description: str, decision: str,
                    content_text: Optional[str] = None,
                    region_code: Optional[str] = None,
                    before_edit: Optional[List[str]] = None,
                    after_edit: Optional[List[str]] = None,
                    reviewer_notes: Optional[str] = None,
                    tags: Optional[List[str]] = None,
                    created_by: str = "system") -> AuditCase:
        case = AuditCase(
            id=str(uuid.uuid4()),
            title=title,
            material_type=material_type,
            content_text=content_text,
            violation_type=violation_type,
            region_code=region_code,
            description=description,
            decision=decision,
            before_edit=before_edit or [],
            after_edit=after_edit or [],
            reviewer_notes=reviewer_notes,
            tags=tags or [],
            created_by=created_by
        )
        self.db.add(case)
        self.db.commit()
        self.db.refresh(case)
        return case

    def search_documents(self, query: str, region_code: Optional[str] = None,
                         doc_type: Optional[str] = None,
                         page: int = 1, page_size: int = 20):
        query_obj = self.db.query(KnowledgeDocument).filter(
            KnowledgeDocument.is_active == True
        )

        if region_code:
            region = self.db.query(Region).filter(Region.code == region_code).first()
            if region:
                query_obj = query_obj.filter(
                    or_(KnowledgeDocument.region_id == region.id,
                        KnowledgeDocument.region_id == None)
                )

        if doc_type:
            try:
                query_obj = query_obj.filter(KnowledgeDocument.doc_type == DocType[doc_type])
            except KeyError:
                pass

        if query:
            query_obj = query_obj.filter(
                or_(
                    KnowledgeDocument.title.ilike(f"%{query}%"),
                    KnowledgeDocument.content_text.ilike(f"%{query}%")
                )
            )

        total = query_obj.count()
        docs = query_obj.order_by(KnowledgeDocument.created_at.desc())\
            .offset((page - 1) * page_size).limit(page_size).all()

        return docs, total

    def list_custom_rules(self, user_id: Optional[str] = None,
                          region_code: Optional[str] = None,
                          page: int = 1, page_size: int = 20):
        query_obj = self.db.query(CustomRule).filter(CustomRule.is_active == True)

        if user_id:
            query_obj = query_obj.filter(CustomRule.user_id == user_id)

        if region_code:
            region = self.db.query(Region).filter(Region.code == region_code).first()
            if region:
                query_obj = query_obj.filter(CustomRule.region_id == region.id)

        total = query_obj.count()
        rules = query_obj.order_by(CustomRule.created_at.desc())\
            .offset((page - 1) * page_size).limit(page_size).all()
        return rules, total

    def list_cases(self, region_code: Optional[str] = None,
                   violation_type: Optional[str] = None,
                   page: int = 1, page_size: int = 20):
        query_obj = self.db.query(AuditCase).filter(AuditCase.is_active == True)

        if region_code:
            query_obj = query_obj.filter(AuditCase.region_code == region_code)
        if violation_type:
            query_obj = query_obj.filter(AuditCase.violation_type == violation_type)

        total = query_obj.count()
        cases = query_obj.order_by(AuditCase.created_at.desc())\
            .offset((page - 1) * page_size).limit(page_size).all()
        return cases, total

    def delete_document(self, doc_id: str) -> bool:
        doc = self.db.query(KnowledgeDocument).filter(KnowledgeDocument.id == doc_id).first()
        if not doc:
            return False
        doc.is_active = False
        self.db.commit()
        return True

    def delete_custom_rule(self, rule_id: str) -> bool:
        rule = self.db.query(CustomRule).filter(CustomRule.id == rule_id).first()
        if not rule:
            return False
        rule.is_active = False
        self.db.commit()
        return True
